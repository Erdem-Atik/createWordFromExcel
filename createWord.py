import pandas as pd
from docx import Document
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Load the Excel file into a pandas dataframe
df = pd.read_excel(r'D:\CODING\PYTHON\BTK_BAZ_İSTASYON\Guv_Sert_Parser\SiteID_AdresKoord.xlsx', header=None)

# Define font style and size
font_style = 'Calibri'
font_size = Pt(9)

# Loop through the rows in the dataframe
for index, row in df.iterrows():
    # Load the template document
    document = Document(r'D:\CODING\PYTHON\BTK_BAZ_İSTASYON\CreateWordFromExcel\application_form.docx')
    
    # Extract the ID, Address, and Coordinates from the row
    id = str(row[0]).replace('.pdf', '').strip()
    address = row[1]
    coordinates = row[2]
    cord_len = (len(coordinates))
    longitutde = coordinates[0:cord_len//2].strip()
    latitude = coordinates[((cord_len//2)+1):cord_len].strip()
        
    # Create a new Word document
    table = document.tables[0]
    cell_1 = table.cell(1, 5)  # add the site ID to template
    cell_2 = table.cell(3, 2)  # add site address to the template
    cell_3 = table.cell(2, 2)  # add site longitutede coordinate to the template
    cell_4 = table.cell(2, 3)  # add site latitude coordinate to the template
    cell_5 = table.cell(6, 4)
    
    paragraph_1 = cell_1.add_paragraph()
    paragraph_2 = cell_2.add_paragraph()
    paragraph_3 = cell_3.add_paragraph()
    paragraph_4 = cell_4.add_paragraph()
    paragraph_5 = cell_5.add_paragraph()

    run_1 = paragraph_1.add_run(id)
    run_1.font.name = font_style
    run_1.font.size = font_size

    run_2 = paragraph_2.add_run(address)
    run_2.font.name = font_style
    run_2.font.size = font_size

    run_3 = paragraph_3.add_run(longitutde)
    run_3.font.name = font_style
    run_3.font.size = font_size

    run_4 = paragraph_4.add_run(latitude)
    run_4.font.name = font_style
    run_4.font.size = font_size

    run_5 = paragraph_5.add_run(address)
    run_5.font.name = font_style
    run_5.font.size = font_size    
    # Save the Word document with the ID and application form as the file name
    filename = f"{id} Başvuru Formu.docx"
    document.save(f"output_folder/{filename}")
